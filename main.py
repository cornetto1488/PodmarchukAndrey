import os
import datetime
import re
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext, CallbackQueryHandler
from docx import Document
from docx.shared import Pt
import sqlite3

def create_table_if_not_exists():
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS user_data (
                    id INTEGER PRIMARY KEY,
                    user_id INTEGER,
                    request_number TEXT,
                    object_id TEXT,
                    desktop_serial TEXT,
                    inv_sticker_number TEXT,
                    vskritie_serial TEXT,
                    nettop_serial TEXT,
                    monitor_serial TEXT,
                    mouse_serial TEXT,
                    keyboard_serial TEXT,
                    ssd_serial TEXT,
                    ram_serial TEXT,
                    employee_name TEXT
                )''')
    conn.commit()
    conn.close()


def get_user_record(chat_id, index, context: CallbackContext) -> str:
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("SELECT * FROM user_data WHERE user_id=?", (chat_id,))
    records = c.fetchall()
    conn.close()

    if records and 0 <= index < len(records):
        record = records[index]
        response = f"Заявка номер: {record[1]}\n"
        response += f"ID объекта: {record[2]}\n"
        response += f"Имя АРМ: {record[3]}\n"
        response += f"Имя сотрудника: {record[13]}\n"
        response += f"Номер инвентаризационной наклейки: {record[4]}\n"
        response += f"Серийный номер наклейки от вскрытия: {record[5]}\n"
        response += f"Серийный номер неттопа: {record[6]}\n"
        response += f"Серийный номер монитора: {record[7]}\n"
        response += f"Серийный номер мышки: {record[8]}\n"
        response += f"Серийный номер клавиатуры: {record[9]}\n"
        response += f"Серийный номер SSD: {record[10]}\n"
        response += f"Серийный номер ОЗУ: {record[11]}\n\nВернуться в главное меню /start\n"
    elif records and index < 0:
        response = "Вы находитесь на первой записи."
    elif records and index >= len(records):
        response = "Вы находитесь на последней записи."
    else:
        response = "У вас нет записей в базе данных."

    return response

def get_record_by_request_number(request_number, record_number, context: CallbackContext):
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("SELECT * FROM user_data WHERE request_number=? LIMIT 1 OFFSET ?", (request_number, record_number - 1))
    record = c.fetchone()
    conn.close()
    return record

def save_data_to_database(context: CallbackContext) -> None:
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    data = context.user_data
    user_id = context.user_data['user_id']  # Получаем user_id из контекста
    c.execute('''INSERT INTO user_data (
                    user_id, request_number, object_id, desktop_serial, inv_sticker_number, vskritie_serial, 
                    nettop_serial, monitor_serial, mouse_serial, keyboard_serial, ssd_serial, 
                    ram_serial, employee_name
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''', (
                    user_id, data['request_number'], data['object_id'], data['desktop_serial'], data['inv_sticker_number'],
                    data['vskritie_serial'], data['nettop_serial'], data['monitor_serial'],
                    data['mouse_serial'], data['keyboard_serial'], data['ssd_serial'],
                    data['ram_serial'], data['employee_name']
                ))
    conn.commit()
    conn.close()

# Функция start, чтобы добавить кнопки
def start(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    context.user_data['user_id'] = user_id  # Устанавливаем user_id в контексте
    # Ваш остальной код обработки команды /start
    keyboard = [
        [
            InlineKeyboardButton("Статистика", callback_data='avr'),
            InlineKeyboardButton("Акт проведения работ", callback_data='work_act'),
            InlineKeyboardButton("АВР", callback_data='rab_act')
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    menu_text = (
        "╔═══════════╗\n"
        "║══ Главное меню ══║\n"
        "║Для удаления записи     ║\n"
        "║/delete_record №Заявки║\n"
        "║═══════════║\n"
        "║ Оставить поле пустым ║\n"
        "║             введите -             ║\n"
        "╚═══════════╝\n"
    )

    update.message.reply_text(menu_text, reply_markup=reply_markup)
    context.user_data['waiting_for'] = 'request_number'


def get_user_records_count(chat_id):
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("SELECT COUNT(id) FROM user_data WHERE user_id=?", (chat_id,))
    user_records_count = c.fetchone()[0]
    conn.close()
    return user_records_count

def get_statistics_from_database(user_id=None):
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("SELECT COUNT(id) FROM user_data")
    total_records = c.fetchone()[0]
    user_records = get_user_records_count(user_id) if user_id else None
    conn.close()

    if user_records is not None:
        return f"Всего записей в базе данных: {total_records}\nВаши записи: {user_records}"
    else:
        return f"Всего записей в базе данных: {total_records}"
def button_click(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    query.answer()
    user_id = update.effective_user.id

    if query.data == 'avr':
        statistics = get_statistics_from_database(user_id)
        user_records = get_user_record(user_id, 0, context)
        context.user_data['current_record_index'] = 0
        keyboard = [
            [
                InlineKeyboardButton("Следующая запись", callback_data='next_record')
            ]
        ] if get_user_records_count(user_id) > 1 else []  # Если записей больше одной, добавляем кнопку для следующей записи
        reply_markup = InlineKeyboardMarkup(keyboard)
        query.edit_message_text(text=f"{statistics}\n\n{user_records}", reply_markup=reply_markup)
    elif query.data == 'work_act':
        query.edit_message_text(text="Введите номер заявки:")
    elif query.data == 'rab_act':
        keyboard = [
            [
                InlineKeyboardButton("Выбрать запись 1", callback_data='choose_record_1'),
                InlineKeyboardButton("Выбрать запись 2", callback_data='choose_record_2'),
                InlineKeyboardButton("Выбрать запись 3", callback_data='choose_record_3')
            ]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        query.edit_message_text(text="Выберите до 3 записей из базы данных.", reply_markup=reply_markup)
    elif query.data.startswith('choose_record_'):
        record_number = int(query.data.split('_')[-1])
        request_number = context.user_data['request_number']
        selected_record = get_record_by_request_number(request_number, record_number, context)
        if selected_record:
            context.user_data[f'selected_record_{record_number}'] = selected_record
            query.edit_message_text(text=f"Запись {record_number} выбрана.")
        else:
            query.edit_message_text(text=f"Запись с номером {request_number} не найдена.")
        if all([f'selected_record_{i}' in context.user_data for i in range(1, 4)]):
            selected_records = [context.user_data[f'selected_record_{i}'] for i in range(1, 4)]
            replace_selected_records_in_document(doc, selected_records)
            update.message.reply_text("Записи успешно заменены в документе.")
            context.user_data.clear()  # Очищаем данные пользователя после выполнения операции
    elif query.data == 'next_record':
        if 'current_record_index' in context.user_data:
            current_index = context.user_data['current_record_index']
            user_records_count = get_user_records_count(user_id)
            if current_index < user_records_count - 1:
                current_index += 1
                user_records = get_user_record(user_id, current_index, context)
                context.user_data['current_record_index'] = current_index
                keyboard = [
                    [
                        InlineKeyboardButton("Предыдущая запись", callback_data='prev_record'),
                        InlineKeyboardButton("Следующая запись", callback_data='next_record')
                    ]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                query.edit_message_text(text=user_records, reply_markup=reply_markup)
            else:
                keyboard = [
                    [
                        InlineKeyboardButton("Предыдущая запись", callback_data='prev_record')
                    ]
                ] if user_records_count > 1 else []  # Если записей больше одной, добавляем кнопку для предыдущей записи
                query.edit_message_text(text="Вы достигли последней записи.\n\nВернуться в главное меню /start", reply_markup=InlineKeyboardMarkup(keyboard))
    elif query.data == 'prev_record':
        if 'current_record_index' in context.user_data:
            current_index = context.user_data['current_record_index']
            if current_index > 0:
                current_index -= 1
                user_records = get_user_record(user_id, current_index, context)
                context.user_data['current_record_index'] = current_index
                keyboard = [
                    [
                        InlineKeyboardButton("Предыдущая запись", callback_data='prev_record'),
                        InlineKeyboardButton("Следующая запись", callback_data='next_record')
                    ]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                query.edit_message_text(text=user_records, reply_markup=reply_markup)
            else:


                keyboard = [
                    [
                        InlineKeyboardButton("Следующая запись", callback_data='next_record')
                    ]
                ] if get_user_records_count(user_id) > 1 else []  # Если записей больше одной, добавляем кнопку для следующей записи
                query.edit_message_text(text="Вы находитесь на первой записи.\n\nВернуться в главное меню /start", reply_markup=InlineKeyboardMarkup(keyboard))


def replace_in_document(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8)
            paragraph.alignment = 0  # left alignment

# Добавим функцию для замены выбранных записей в документе
def replace_selected_records_in_document(doc, selected_records):
    for i, record_index in enumerate(selected_records, start=1):
        old_text = f'{{{{zapis{i}}}}}'
        new_text = get_user_record(update.effective_user.id, record_index, context)
        replace_in_document(doc, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(8)
                        paragraph.alignment = 0  # left alignment

    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)
                    paragraph.alignment = 0  # left alignment

        if section.footer:
            for paragraph in section.footer.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)
                    paragraph.alignment = 0  # left alignment
def replace_datetime_in_document(doc, datetime_text):
    current_date = datetime.datetime.now().strftime("%d.%m.%Y")
    for paragraph in doc.paragraphs:
        if datetime_text in paragraph.text:
            paragraph.text = paragraph.text.replace(datetime_text, current_date)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8)
            paragraph.alignment = 0  # left alignment

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if datetime_text in cell.text:
                    cell.text = cell.text.replace(datetime_text, current_date)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(8)
                        paragraph.alignment = 0  # left alignment

    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                if datetime_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(datetime_text, current_date)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)
                    paragraph.alignment = 0  # left alignment

        if section.footer:
            for paragraph in section.footer.paragraphs:
                if datetime_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(datetime_text, current_date)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)
                    paragraph.alignment = 0  # left alignment

def format_and_display_data(context: CallbackContext) -> str:
    data = context.user_data
    formatted_data = f"Номер заявки: {data['request_number']}\n"
    formatted_data += f"ID объекта: {data['object_id']}\n"
    formatted_data += f"Имя АРМ: {data['desktop_serial']}\n"
    formatted_data += f"Имя сотрудника: {data['employee_name']}\n"
    formatted_data += f"Номер инвентаризационной наклейки: {data['inv_sticker_number']}\n"
    formatted_data += f"Серийный номер наклейки от вскрытия: {data['vskritie_serial']}\n"
    formatted_data += f"Серийный номер неттопа: {data['nettop_serial']}\n"
    formatted_data += f"Серийный номер монитора: {data['monitor_serial']}\n"
    formatted_data += f"Серийный номер мышки: {data['mouse_serial']}\n"
    formatted_data += f"Серийный номер клавиатуры: {data['keyboard_serial']}\n"
    formatted_data += f"Серийный номер SSD: {data['ssd_serial']}\n"
    formatted_data += f"Серийный номер ОЗУ: {data['ram_serial']}\n\nВернуться в главное меню /start\n"
    return formatted_data

def echo(update: Update, context: CallbackContext) -> None:
    if 'waiting_for' in context.user_data:
        waiting_for = context.user_data['waiting_for']
        user_input = update.message.text
        if waiting_for == 'request_number':
            context.user_data['request_number'] = user_input
            update.message.reply_text('Введите ID объекта:')
            context.user_data['waiting_for'] = 'object_id'
        elif waiting_for == 'object_id':
            context.user_data['object_id'] = user_input
            update.message.reply_text('Введите имя сотрудника:')
            context.user_data['waiting_for'] = 'employee_name'  # Изменение состояния ожидания
        elif waiting_for == 'employee_name':
            context.user_data['employee_name'] = user_input  # Сохранение имени сотрудника
            update.message.reply_text('Введите имя АРМ:')
            context.user_data['waiting_for'] = 'desktop_serial'
        elif waiting_for == 'desktop_serial':
            context.user_data['desktop_serial'] = user_input
            update.message.reply_text('Введите номер инвентаризационной наклейки:')
            context.user_data['waiting_for'] = 'inv_sticker_number'
        elif waiting_for == 'inv_sticker_number':
            context.user_data['inv_sticker_number'] = user_input
            update.message.reply_text('Введите серийный номер наклейки от вскрытия:')
            context.user_data['waiting_for'] = 'vskritie_serial'
        elif waiting_for == 'vskritie_serial':
            context.user_data['vskritie_serial'] = user_input
            update.message.reply_text('Введите серийный номер неттопа:')
            context.user_data['waiting_for'] = 'nettop_serial'
        elif waiting_for == 'nettop_serial':
            context.user_data['nettop_serial'] = user_input
            update.message.reply_text('Введите серийный номер монитора:')
            context.user_data['waiting_for'] = 'monitor_serial'
        elif waiting_for == 'monitor_serial':
            context.user_data['monitor_serial'] = user_input
            update.message.reply_text('Введите серийный номер мышки:')
            context.user_data['waiting_for'] = 'mouse_serial'
        elif waiting_for == 'mouse_serial':
            context.user_data['mouse_serial'] = user_input
            update.message.reply_text('Введите серийный номер клавиатуры:')
            context.user_data['waiting_for'] = 'keyboard_serial'
        elif waiting_for == 'keyboard_serial':
            context.user_data['keyboard_serial'] = user_input
            update.message.reply_text('Введите серийный номер SSD:')
            context.user_data['waiting_for'] = 'ssd_serial'
        elif waiting_for == 'ssd_serial':
            context.user_data['ssd_serial'] = user_input
            update.message.reply_text('Введите серийный номер ОЗУ:')
            context.user_data['waiting_for'] = 'ram_serial'
        elif waiting_for == 'ram_serial':
            context.user_data['ram_serial'] = user_input
            formatted_data = format_and_display_data(context)
            update.message.reply_text(f'Вы ввели следующие данные:\n{formatted_data}')
            save_data_to_database(context)
            save_data_to_document(update, context)
            context.user_data.clear()
        else:
            update.message.reply_text('Что-то пошло не так, попробуйте еще раз.')
    else:
        update.message.reply_text('Привет! Пожалуйста, начните с команды /start.')
def delete_record(update: Update, context: CallbackContext) -> None:
    if len(context.args) == 0:
        update.message.reply_text("Пожалуйста, укажите номер заявки для удаления.")
        return

    request_number = context.args[0]
    user_id = update.message.from_user.id

    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("DELETE FROM user_data WHERE request_number=? AND user_id=?", (request_number, user_id))
    conn.commit()
    conn.close()

    update.message.reply_text(f"Запись с номером заявки {request_number} успешно удалена.")
def get_nameobject_from_database(object_id):
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("SELECT nameobject FROM object_information WHERE id=?", (object_id,))
    row = c.fetchone()
    conn.close()
    return row[0] if row else None
def get_address_from_database(object_id):
    conn = sqlite3.connect('user_data.db')
    c = conn.cursor()
    c.execute("SELECT adressobject FROM object_information WHERE id=?", (object_id,))
    row = c.fetchone()
    conn.close()
    return row[0] if row else None
def save_data_to_document(update: Update, context: CallbackContext) -> None:
    request_number = context.user_data['request_number']
    object_id = context.user_data['object_id']
    employee_name = context.user_data.get('employee_name', '-')
    desktop_serial = context.user_data['desktop_serial']
    monitor_serial = context.user_data['monitor_serial']
    vskritie_serial = context.user_data['vskritie_serial']
    inv_sticker_number = context.user_data['inv_sticker_number']
    nettop_serial = context.user_data['nettop_serial']
    mouse_serial = context.user_data['mouse_serial']
    keyboard_serial = context.user_data['keyboard_serial']
    ssd_serial = context.user_data['ssd_serial']
    ram_serial = context.user_data['ram_serial']
    template_path = '1.docx'
    doc = Document(template_path)
    object_address = get_address_from_database(object_id)
    if object_address:
        replace_in_document(doc, '{{addressobject}}', object_address)
    else:
        replace_in_document(doc, '{{addressobject}}', "-")  # Если адрес не найден, заменяем на "-"
    object_information = get_nameobject_from_database(object_id)
    if object_information:
        replace_in_document(doc, '{{nameobject}}', object_information)
    else:
        replace_in_document(doc, '{{nameobject}}', "-")  # Если информация не найдена, заменяем на "-"

    replace_in_document(doc, '{{zayavka}}', request_number)
    replace_in_document(doc, '{{id_object}}', object_id)  # Заменил здесь
    replace_in_document(doc, '{{doc_number}}', desktop_serial)
    replace_in_document(doc, '{{sn_monitor}}', monitor_serial)
    replace_in_document(doc, '{{vskritie}}', vskritie_serial)
    replace_in_document(doc, '{{sn_nettop}}', nettop_serial)
    replace_in_document(doc, '{{sn_mouse}}', mouse_serial)
    replace_in_document(doc, '{{sn_keyboard}}', keyboard_serial)
    replace_in_document(doc, '{{sn_ssd}}', ssd_serial)
    replace_in_document(doc, '{{rtk_numb}}', inv_sticker_number)
    replace_in_document(doc, '{{username}}', employee_name)
    replace_in_document(doc, '{{sn_ozy}}', ram_serial)
    replace_datetime_in_document(doc, '{{datetime}}')
    directory = 'C:\\Workplace\\project\\akti'
    filename_base = f'{directory}/PPV_{request_number}.docx'
    filename = filename_base
    counter = 1
    while os.path.exists(filename):
        filename = f"{filename_base}_{counter}.docx"
        counter += 1
    doc.save(filename)
    update.message.reply_document(document=open(filename, 'rb'))
def main() -> None:
    create_table_if_not_exists()
    TOKEN = '7108768960:AAEKo0NbqmfLFKlF2Pn-1gMByAh8WxwxtkY'
    updater = Updater(TOKEN)
    dispatcher = updater.dispatcher
    dispatcher.add_handler(CommandHandler("start", start))
    dispatcher.add_handler(CommandHandler("my_records", get_user_record))
    dispatcher.add_handler(CallbackQueryHandler(button_click))
    dispatcher.add_handler(MessageHandler(Filters.text & ~Filters.command, echo))
    dispatcher.add_handler(CommandHandler("delete_record", delete_record))
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()
